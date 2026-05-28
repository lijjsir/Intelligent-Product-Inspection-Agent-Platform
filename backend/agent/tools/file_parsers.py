from __future__ import annotations

import csv
import json
import re
import zipfile
from io import BytesIO, StringIO
from pathlib import Path
from xml.etree import ElementTree


def parse_pdf_bytes(content: bytes) -> dict:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - runtime dependency
        raise ValueError(f"pdf parsing dependency unavailable: {exc}") from exc

    reader = PdfReader(BytesIO(content))
    page_texts = [(page.extract_text() or "").strip() for page in reader.pages]
    pages = [
        {
            "page_no": index,
            "text": page_text,
            "blocks": ([{"type": "text", "text": page_text}] if page_text else []),
            "fonts": [],
            "figures": [],
            "tables": [],
        }
        for index, page_text in enumerate(page_texts, start=1)
    ]
    text = "\n".join(page for page in page_texts if page)
    return {
        "kind": "pdf",
        "page_count": len(reader.pages),
        "pages": pages,
        "headings": _extract_pdf_headings(text),
        "references": _extract_reference_lines(text),
        "layout": {
            "analysis": "text_only",
            "parser": "pypdf",
            "page_count": len(reader.pages),
        },
        "text": text,
    }


def parse_docx_bytes(content: bytes) -> dict:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as exc:  # pragma: no cover - runtime dependency
        raise ValueError(f"docx parsing dependency unavailable: {exc}") from exc

    doc = Document(BytesIO(content))
    paragraphs: list[dict] = []
    headings: list[dict] = []
    figure_titles: list[str] = []

    alignment_map = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribute",
    }

    current_section_title = ""
    current_section_level = 0
    current_section_index = 0
    current_paragraph_no = 0

    for index, paragraph in enumerate(doc.paragraphs):
        text = (paragraph.text or "").strip()
        style_name = str(getattr(paragraph.style, "name", "") or "")
        heading_level = _heading_level(style_name)
        run_fonts = []
        run_sizes = []
        bold = False
        for run in paragraph.runs:
            if run.bold:
                bold = True
            if run.font.name:
                run_fonts.append(str(run.font.name))
            if run.font.size:
                run_sizes.append(round(run.font.size.pt, 2))
        para_format = paragraph.paragraph_format
        line_spacing = _to_float(getattr(para_format, "line_spacing", None))
        space_before = _to_pt(getattr(para_format, "space_before", None))
        space_after = _to_pt(getattr(para_format, "space_after", None))
        first_line_indent = _to_pt(getattr(para_format, "first_line_indent", None))
        alignment = alignment_map.get(paragraph.alignment, str(paragraph.alignment).lower() if paragraph.alignment is not None else "")
        if heading_level:
            current_section_title = text
            current_section_level = heading_level
            current_section_index += 1
            current_paragraph_no = 0
        elif text:
            current_paragraph_no += 1

        item = {
            "index": index,
            "text": text,
            "style_name": style_name,
            "heading_level": heading_level,
            "font_name": run_fonts[0] if run_fonts else "",
            "font_size_pt": run_sizes[0] if run_sizes else None,
            "bold": bold,
            "alignment": alignment,
            "line_spacing": line_spacing,
            "space_before_pt": space_before,
            "space_after_pt": space_after,
            "first_line_indent_pt": first_line_indent,
            "section_title": current_section_title,
            "section_level": current_section_level,
            "section_index": current_section_index,
            "paragraph_no": current_paragraph_no if text and not heading_level else 0,
        }
        paragraphs.append(item)
        if heading_level:
            headings.append(
                {
                    "text": text,
                    "level": heading_level,
                    "paragraph_index": index,
                    "section_index": current_section_index,
                    "font_name": item["font_name"],
                    "font_size_pt": item["font_size_pt"],
                    "bold": bold,
                }
            )
        if text and re.match(r"^(图|Figure)\s*\d+", text, re.I):
            figure_titles.append(text)

    sections = []
    for section in doc.sections:
        sections.append(
            {
                "start_type": str(getattr(section.start_type, "name", "") or ""),
                "page_width_cm": _to_cm(getattr(section, "page_width", None)),
                "page_height_cm": _to_cm(getattr(section, "page_height", None)),
                "top_margin_cm": _to_cm(getattr(section, "top_margin", None)),
                "bottom_margin_cm": _to_cm(getattr(section, "bottom_margin", None)),
                "left_margin_cm": _to_cm(getattr(section, "left_margin", None)),
                "right_margin_cm": _to_cm(getattr(section, "right_margin", None)),
                "header_distance_cm": _to_cm(getattr(section, "header_distance", None)),
                "footer_distance_cm": _to_cm(getattr(section, "footer_distance", None)),
                "header_text": "\n".join(
                    (paragraph.text or "").strip()
                    for paragraph in section.header.paragraphs
                    if (paragraph.text or "").strip()
                ),
                "footer_text": "\n".join(
                    (paragraph.text or "").strip()
                    for paragraph in section.footer.paragraphs
                    if (paragraph.text or "").strip()
                ),
            }
        )

    text_lines = [item["text"] for item in paragraphs if item["text"]]
    return {
        "kind": "docx",
        "paragraphs": paragraphs,
        "headings": headings,
        "figure_titles": figure_titles,
        "section_count": len(doc.sections),
        "page_layout": sections[0] if sections else {},
        "sections": sections,
        "text": "\n".join(text_lines),
    }


def parse_tex_bytes(content: bytes) -> dict:
    text = content.decode("utf-8", errors="ignore")
    lines = text.splitlines()
    commands = {
        "documentclass": _extract_tex_command(text, "documentclass"),
        "title": _extract_tex_command(text, "title"),
        "author": _extract_tex_command(text, "author"),
        "abstract": _extract_tex_environment(text, "abstract"),
        "bibliography": _extract_tex_command(text, "bibliography") or _extract_tex_environment(text, "thebibliography"),
    }
    packages = re.findall(r"\\usepackage(?:\[[^\]]*\])?\{([^}]+)\}", text)
    sections = []
    paragraphs = []
    figure_titles = re.findall(r"\\caption\{([^}]*)\}", text)
    for lineno, line in enumerate(lines, start=1):
        stripped = line.strip()
        if not stripped or stripped.startswith("%"):
            continue
        section_title = sections[-1]["title"] if sections else ""
        section_index = len(sections)
        section_match = re.match(r"\\(part|chapter|section|subsection|subsubsection)\*?\{(.+?)\}", stripped)
        if section_match:
            sections.append(
                {
                    "command": section_match.group(1),
                    "title": section_match.group(2).strip(),
                    "line": lineno,
                }
            )
            section_title = section_match.group(2).strip()
            section_index = len(sections)
        paragraphs.append(
            {
                "index": len(paragraphs),
                "line": lineno,
                "text": stripped,
                "heading_level": 0,
                "section_title": section_title,
                "section_index": section_index,
            }
        )
    return {
        "kind": "tex",
        "text": text,
        "commands": commands,
        "packages": packages,
        "sections": sections,
        "paragraphs": paragraphs,
        "headings": [
            {
                "text": item["title"],
                "level": {"part": 1, "chapter": 1, "section": 1, "subsection": 2, "subsubsection": 3}.get(item["command"], 0),
                "paragraph_index": item["line"],
            }
            for item in sections
        ],
        "figure_titles": figure_titles,
        "figure_count": len(re.findall(r"\\begin\{figure\}", text)),
        "table_count": len(re.findall(r"\\begin\{table\}", text)),
    }


def parse_csv_bytes(content: bytes) -> dict:
    text = content.decode("utf-8", errors="ignore")
    reader = csv.reader(StringIO(text))
    rows = list(reader)
    headers = rows[0] if rows else []
    return {
        "kind": "csv",
        "headers": headers,
        "rows": rows[1:51] if len(rows) > 1 else [],
        "text": text,
    }


def parse_json_bytes(content: bytes) -> dict:
    text = content.decode("utf-8", errors="ignore")
    payload = json.loads(text or "{}")
    return {
        "kind": "json",
        "payload": payload,
        "text": json.dumps(payload, ensure_ascii=False, indent=2),
    }


def parse_xlsx_bytes(content: bytes) -> dict:
    with zipfile.ZipFile(BytesIO(content)) as archive:
        shared_strings = []
        if "xl/sharedStrings.xml" in archive.namelist():
            xml = archive.read("xl/sharedStrings.xml").decode("utf-8", errors="ignore")
            root = ElementTree.fromstring(xml)
            shared_strings = ["".join(node.itertext()) for node in root]
        sheet_entries = []
        for name in archive.namelist():
            if not name.startswith("xl/worksheets/sheet") or not name.endswith(".xml"):
                continue
            xml = archive.read(name).decode("utf-8", errors="ignore")
            root = ElementTree.fromstring(xml)
            rows = []
            for row in root.iter():
                if row.tag.endswith("row"):
                    cells = []
                    for cell in row:
                        if not cell.tag.endswith("c"):
                            continue
                        value = ""
                        cell_type = cell.attrib.get("t")
                        v_node = next((child for child in cell if child.tag.endswith("v")), None)
                        if v_node is not None and v_node.text is not None:
                            value = v_node.text
                            if cell_type == "s":
                                try:
                                    value = shared_strings[int(value)]
                                except Exception:
                                    pass
                        cells.append(value)
                    if cells:
                        rows.append(cells)
            sheet_entries.append({"sheet": Path(name).stem, "rows": rows[:51]})
    text_lines = []
    for sheet in sheet_entries:
        text_lines.append(f"[{sheet['sheet']}]")
        for row in sheet["rows"][:20]:
            text_lines.append(" | ".join(str(item) for item in row))
    return {
        "kind": "xlsx",
        "sheets": sheet_entries,
        "text": "\n".join(text_lines),
    }


def parse_text_bytes(content: bytes) -> dict:
    text = content.decode("utf-8", errors="ignore")
    return {"kind": "text", "text": text}


def parse_file_content(file_name: str, content: bytes) -> dict:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".pdf":
        return parse_pdf_bytes(content)
    if suffix == ".docx":
        return parse_docx_bytes(content)
    if suffix == ".tex":
        return parse_tex_bytes(content)
    if suffix == ".csv":
        return parse_csv_bytes(content)
    if suffix == ".json":
        return parse_json_bytes(content)
    if suffix == ".xlsx":
        return parse_xlsx_bytes(content)
    return parse_text_bytes(content)


def _heading_level(style_name: str) -> int:
    match = re.search(r"heading\s*(\d+)", style_name, re.I)
    if match:
        return int(match.group(1))
    if "标题" in style_name:
        zh_match = re.search(r"标题\s*(\d+)", style_name)
        if zh_match:
            return int(zh_match.group(1))
        return 1
    return 0


def _extract_pdf_headings(text: str) -> list[dict]:
    headings: list[dict] = []
    for index, line in enumerate(str(text or "").splitlines()):
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"^\d+(?:\.\d+)*\s+\S+", stripped) or stripped.lower() in {
            "abstract",
            "introduction",
            "references",
        } or stripped in {"摘要", "关键词", "目录", "参考文献", "致谢"}:
            headings.append(
                {
                    "text": stripped,
                    "level": _pdf_heading_level(stripped),
                    "paragraph_index": index,
                }
            )
    return headings


def _pdf_heading_level(line: str) -> int:
    match = re.match(r"^(\d+(?:\.\d+)*)\s+\S+", line)
    if not match:
        return 1
    return min(match.group(1).count(".") + 1, 6)


def _extract_reference_lines(text: str) -> list[str]:
    references: list[str] = []
    in_references = False
    for line in str(text or "").splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.lower() in {"references", "bibliography"} or stripped == "参考文献":
            in_references = True
            continue
        if in_references and (
            re.match(r"^\[\d+\]\s*\S+", stripped)
            or re.match(r"^\d+\.\s*\S+", stripped)
        ):
            references.append(stripped)
    return references


def _to_float(value) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    try:
        return float(value)
    except Exception:
        return None


def _to_pt(value) -> float | None:
    if value is None:
        return None
    try:
        return round(float(value.pt), 2)
    except Exception:
        return None


def _to_cm(value) -> float | None:
    if value is None:
        return None
    try:
        return round(float(value.cm), 2)
    except Exception:
        return None


def _extract_tex_command(text: str, name: str) -> str:
    match = re.search(rf"\\{name}(?:\[[^\]]*\])?\{{([^}}]*)\}}", text, re.S)
    return match.group(1).strip() if match else ""


def _extract_tex_environment(text: str, name: str) -> str:
    match = re.search(rf"\\begin\{{{name}\}}(.*?)\\end\{{{name}\}}", text, re.S)
    return re.sub(r"\s+", " ", match.group(1)).strip() if match else ""
