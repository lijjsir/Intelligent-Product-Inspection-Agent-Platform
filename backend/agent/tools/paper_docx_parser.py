"""Enhanced DOCX parser with style inheritance resolution, table parsing, and citation extraction."""
from __future__ import annotations

import re
import zipfile
from io import BytesIO
from pathlib import PurePosixPath
from typing import Any


def parse_docx_enhanced(content: bytes) -> dict[str, Any]:
    """Parse DOCX with style inheritance resolution and enhanced extraction."""
    from docx import Document

    doc = Document(BytesIO(content))
    style_cache = _build_style_cache(doc)
    xml_bundle = _parse_docx_xml_bundle(content)
    xml_paragraphs = list(xml_bundle.get("paragraphs") or [])
    section_breaks = list(xml_bundle.get("section_breaks") or [])
    parser_limitations: list[str] = []
    if xml_bundle.get("error"):
        parser_limitations.append("底层样式解析未完成，部分版式检查已降级。")

    paragraphs = []
    headings = []
    figure_titles = []

    current_section_title = ""
    current_section_level = 0
    current_section_index = 0
    current_paragraph_no = 0
    current_word_section_index = 0
    for index, paragraph in enumerate(doc.paragraphs):
        while current_word_section_index + 1 < len(section_breaks) and index >= int(section_breaks[current_word_section_index + 1]):
            current_word_section_index += 1
        text = (paragraph.text or "").strip()
        style_name = str(getattr(paragraph.style, "name", "") or "")
        style_id = str(getattr(paragraph.style, "style_id", "") or "")
        heading_level = _heading_level(style_name)
        xml_paragraph = xml_paragraphs[index] if index < len(xml_paragraphs) else {}

        font_name_raw, font_name_resolved, style_source = _resolve_font(paragraph, style_cache, style_name, style_id)
        font_size_raw, font_size_resolved, size_source = _resolve_font_size(paragraph, style_cache, style_name, style_id)

        para_format = paragraph.paragraph_format
        line_spacing, line_spacing_source = _resolve_paragraph_metric(
            direct_value=getattr(para_format, "line_spacing", None),
            style_cache=style_cache,
            style_name=style_name,
            style_id=style_id,
            key="line_spacing",
            converter=_to_float,
        )
        space_before, space_before_source = _resolve_paragraph_metric(
            direct_value=getattr(para_format, "space_before", None),
            style_cache=style_cache,
            style_name=style_name,
            style_id=style_id,
            key="space_before_pt",
            converter=_to_pt,
        )
        space_after, space_after_source = _resolve_paragraph_metric(
            direct_value=getattr(para_format, "space_after", None),
            style_cache=style_cache,
            style_name=style_name,
            style_id=style_id,
            key="space_after_pt",
            converter=_to_pt,
        )
        first_line_indent, first_line_indent_source = _resolve_paragraph_metric(
            direct_value=getattr(para_format, "first_line_indent", None),
            style_cache=style_cache,
            style_name=style_name,
            style_id=style_id,
            key="first_line_indent_pt",
            converter=_to_pt,
        )
        alignment, alignment_source = _resolve_paragraph_metric(
            direct_value=paragraph.alignment,
            style_cache=style_cache,
            style_name=style_name,
            style_id=style_id,
            key="alignment",
            converter=_alignment_token,
        )
        bold = any(bool(getattr(run, "bold", False)) for run in paragraph.runs)

        if heading_level:
            current_section_title = text
            current_section_level = heading_level
            current_section_index += 1
            current_paragraph_no = 0
        elif text:
            current_paragraph_no += 1

        item = {
            "index": index,
            "paragraph_index": index,
            "text": text,
            "style_name": style_name,
            "style_id": style_id or str(xml_paragraph.get("style_id") or ""),
            "heading_level": heading_level,
            "font_name_raw": font_name_raw,
            "font_name_resolved": font_name_resolved,
            "font_name": font_name_resolved or font_name_raw,
            "font_style_source": style_source,
            "font_size_raw": font_size_raw,
            "font_size_resolved": font_size_resolved,
            "font_size_pt": font_size_resolved or font_size_raw,
            "font_size_source": size_source,
            "line_spacing": line_spacing,
            "line_spacing_source": line_spacing_source,
            "space_before_pt": space_before,
            "space_before_source": space_before_source,
            "space_after_pt": space_after,
            "space_after_source": space_after_source,
            "first_line_indent_pt": first_line_indent,
            "first_line_indent_source": first_line_indent_source,
            "alignment": alignment,
            "alignment_source": alignment_source,
            "bold": bold,
            "section_title": current_section_title,
            "section_level": current_section_level,
            "section_index": current_section_index,
            "word_section_index": current_word_section_index,
            "paragraph_no": current_paragraph_no if text and not heading_level else 0,
            "xml_path": xml_paragraph.get("xml_path"),
            "text_runs": list(xml_paragraph.get("text_runs") or []),
            "numbering": dict(xml_paragraph.get("numbering") or {}),
            "page_break_before": bool(xml_paragraph.get("page_break_before")),
            "has_page_break": bool(xml_paragraph.get("has_page_break")),
            "footnote_refs": list(xml_paragraph.get("footnote_refs") or []),
            "endnote_refs": list(xml_paragraph.get("endnote_refs") or []),
            "comment_refs": list(xml_paragraph.get("comment_refs") or []),
            "field_codes": list(xml_paragraph.get("field_codes") or []),
        }
        paragraphs.append(item)

        if heading_level:
            headings.append({
                "text": text,
                "level": heading_level,
                "paragraph_index": index,
                "section_index": current_section_index,
                "font_name": font_name_resolved or font_name_raw or "",
                "font_size_pt": font_size_resolved or font_size_raw,
            })

        if text and re.match(r"^(图|Figure)\s*\d+(?:[-.]\d+)*", text, re.I):
            figure_titles.append(text)

    tables = _parse_tables(doc)
    sections_data = _parse_sections_enhanced(doc)
    text_content = "\n".join(p["text"] for p in paragraphs if p["text"])
    citations = _extract_citations(text_content)
    table_titles = _extract_titles(text_content, r"^(?:表|Table)\s*\d+(?:[-.]\d+)*[^\n]*")
    formula_numbers = _extract_formula_numbers(paragraphs, text_content)
    toc_entries = _extract_toc_entries(paragraphs)
    word_metadata = _parse_docx_package_metadata(content)

    return {
        "kind": "docx",
        "paragraphs": paragraphs,
        "headings": headings,
        "figure_titles": figure_titles,
        "table_titles": table_titles,
        "formula_numbers": formula_numbers,
        "toc_entries": toc_entries,
        "tables": tables,
        "sections": sections_data,
        "section_count": len(doc.sections),
        "page_layout": sections_data[0] if sections_data else {},
        "citations": citations,
        "references": _extract_reference_lines(text_content),
        "word_metadata": word_metadata,
        "parser_limitations": parser_limitations,
        "ooxml": dict(xml_bundle.get("meta") or {}),
        "text": text_content,
    }


def _build_style_cache(doc) -> dict[str, dict[str, Any]]:
    by_name: dict[str, dict[str, Any]] = {}
    by_id: dict[str, dict[str, Any]] = {}
    resolved: dict[str, dict[str, Any]] = {}

    def resolve_style(style) -> dict[str, Any]:
        if style is None:
            return {}
        cache_key = str(getattr(style, "style_id", "") or getattr(style, "name", "") or id(style))
        if cache_key in resolved:
            return resolved[cache_key]
        base_info = resolve_style(getattr(style, "base_style", None))
        current = {
            "font_name": _style_font_name(style),
            "font_size_pt": _to_pt(_safe_getattr(getattr(style, "font", None), "size")),
            "line_spacing": _to_float(_safe_getattr(getattr(style, "paragraph_format", None), "line_spacing")),
            "space_before_pt": _to_pt(_safe_getattr(getattr(style, "paragraph_format", None), "space_before")),
            "space_after_pt": _to_pt(_safe_getattr(getattr(style, "paragraph_format", None), "space_after")),
            "first_line_indent_pt": _to_pt(_safe_getattr(getattr(style, "paragraph_format", None), "first_line_indent")),
            "alignment": _alignment_token(_safe_getattr(getattr(style, "paragraph_format", None), "alignment")),
        }
        merged = dict(base_info)
        for key, value in current.items():
            if value not in (None, ""):
                merged[key] = value
            else:
                merged.setdefault(key, value)
        resolved[cache_key] = merged
        return merged

    for style in doc.styles:
        info = resolve_style(style)
        style_name = str(getattr(style, "name", "") or "")
        style_id = str(getattr(style, "style_id", "") or "")
        if style_name:
            by_name[style_name] = info
        if style_id:
            by_id[style_id] = info
    return {"by_name": by_name, "by_id": by_id}


def _resolve_font(paragraph, style_cache: dict[str, Any], style_name: str, style_id: str) -> tuple:
    for run in paragraph.runs:
        if run.font.name:
            return run.font.name, run.font.name, "run"

    fn = _style_cache_value(style_cache, style_name=style_name, style_id=style_id, key="font_name")
    if fn:
        return None, fn, "paragraph_style"

    return None, None, "unresolved"


def _resolve_font_size(paragraph, style_cache: dict[str, Any], style_name: str, style_id: str) -> tuple:
    for run in paragraph.runs:
        if run.font.size:
            try:
                return run.font.size.pt, round(run.font.size.pt, 2), "run"
            except Exception:
                pass

    fs = _style_cache_value(style_cache, style_name=style_name, style_id=style_id, key="font_size_pt")
    if isinstance(fs, (int, float)):
        return None, round(float(fs), 2), "paragraph_style"

    return None, None, "unresolved"


def _resolve_paragraph_metric(
    *,
    direct_value,
    style_cache: dict[str, Any],
    style_name: str,
    style_id: str,
    key: str,
    converter,
) -> tuple[Any, str]:
    direct_resolved = converter(direct_value)
    if direct_resolved not in (None, ""):
        return direct_resolved, "paragraph"
    style_value = _style_cache_value(style_cache, style_name=style_name, style_id=style_id, key=key)
    if style_value not in (None, ""):
        return style_value, "paragraph_style"
    return direct_resolved, "unresolved"


def _style_cache_value(style_cache: dict[str, Any], *, style_name: str, style_id: str, key: str) -> Any:
    by_id = dict(style_cache.get("by_id") or {})
    by_name = dict(style_cache.get("by_name") or {})
    candidates = []
    if style_id:
        candidates.append(by_id.get(style_id))
    if style_name:
        candidates.append(by_name.get(style_name))
    candidates.append(by_name.get("Normal"))
    for item in candidates:
        if not isinstance(item, dict):
            continue
        value = item.get(key)
        if value not in (None, ""):
            return value
    return None


def _style_font_name(style) -> str | None:
    try:
        return getattr(getattr(style, "font", None), "name", None)
    except Exception:
        return None


def _safe_getattr(value, name: str):
    try:
        return getattr(value, name, None)
    except Exception:
        return None


def _alignment_token(value) -> str:
    raw = str(getattr(value, "name", value) or "").strip().lower()
    if not raw:
        return ""
    if raw.startswith("justify"):
        return "justify"
    if raw.startswith("center") or raw.startswith("centre"):
        return "center"
    if raw.startswith("right"):
        return "right"
    if raw.startswith("left"):
        return "left"
    if raw.startswith("distribute") or raw.startswith("distributed"):
        return "distribute"
    return raw


def _parse_tables(doc) -> list[dict[str, Any]]:
    tables = []
    for t_idx, table in enumerate(doc.tables):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append({"index": t_idx, "rows": rows, "row_count": len(rows)})
    return tables


def _parse_sections_enhanced(doc) -> list[dict[str, Any]]:
    sections = []
    for section in doc.sections:
        page_width = _to_cm(getattr(section, "page_width", None))
        page_height = _to_cm(getattr(section, "page_height", None))
        sections.append({
            "start_type": str(getattr(getattr(section, "start_type", None), "name", "") or getattr(section, "start_type", "") or ""),
            "page_width_cm": page_width,
            "page_height_cm": page_height,
            "orientation": (
                "landscape"
                if isinstance(page_width, (int, float)) and isinstance(page_height, (int, float)) and page_width > page_height
                else "portrait"
            ),
            "top_margin_cm": _to_cm(getattr(section, "top_margin", None)),
            "bottom_margin_cm": _to_cm(getattr(section, "bottom_margin", None)),
            "left_margin_cm": _to_cm(getattr(section, "left_margin", None)),
            "right_margin_cm": _to_cm(getattr(section, "right_margin", None)),
            "gutter_cm": _to_cm(getattr(section, "gutter", None)),
            "header_distance_cm": _to_cm(getattr(section, "header_distance", None)),
            "footer_distance_cm": _to_cm(getattr(section, "footer_distance", None)),
            "header_text": "\n".join(
                (p.text or "").strip() for p in section.header.paragraphs if (p.text or "").strip()
            ),
            "footer_text": "\n".join(
                (p.text or "").strip() for p in section.footer.paragraphs if (p.text or "").strip()
            ),
        })
    return sections


def _extract_titles(text: str, pattern: str) -> list[str]:
    return [
        match.group(0).strip()
        for match in re.finditer(pattern, text, flags=re.I | re.M)
        if match.group(0).strip()
    ]


def _extract_formula_numbers(paragraphs: list[dict[str, Any]], text: str) -> list[str]:
    numbers: list[str] = []
    for paragraph in paragraphs:
        paragraph_text = str(paragraph.get("text") or "")
        style_name = str(paragraph.get("style_name") or "")
        if "公式" not in style_name and not re.search(r"[=∑√≤≥±×÷]", paragraph_text):
            continue
        for match in re.finditer(r"[（(]\s*(\d+(?:[-.]\d+)*)\s*[）)]", paragraph_text):
            number = match.group(1)
            if _plausible_formula_number(number) and number not in numbers:
                numbers.append(number)
    return numbers


def _plausible_formula_number(number: str) -> bool:
    parts = re.split(r"[-.]", str(number or ""))
    if not 1 <= len(parts) <= 3:
        return False
    try:
        values = [int(part) for part in parts]
    except ValueError:
        return False
    return values[0] <= 20 and all(0 <= item <= 99 for item in values)


def _extract_toc_entries(paragraphs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for paragraph in paragraphs:
        style_name = str(paragraph.get("style_name") or "").lower()
        text = str(paragraph.get("text") or "").strip()
        if not text:
            continue
        if not (style_name.startswith("toc") or "table of figures" in style_name):
            continue
        title = text
        page = ""
        if "\t" in text:
            title, page = text.rsplit("\t", 1)
        entries.append({
            "title": title.strip(),
            "page": page.strip(),
            "style_name": paragraph.get("style_name"),
            "level": _toc_level(style_name),
            "paragraph_index": paragraph.get("index"),
        })
    return entries


def _toc_level(style_name: str) -> int:
    match = re.search(r"toc\s*(\d+)", style_name)
    if match:
        return int(match.group(1))
    return 1


def _extract_reference_lines(text: str) -> list[str]:
    lines = []
    in_refs = False
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        # Only trigger on exact section header, not on entries that contain the word
        if in_refs:
            if re.match(r"^\[\d+\]", line):
                lines.append(line)
            # Stop at next major section heading
            elif re.match(r"^(?:致谢|附录|攻读|发表|在学|致 謝)", line):
                break
        else:
            if line in {"参考文献", "References", "Bibliography", "REFERENCES"}:
                in_refs = True
    return lines


def _parse_docx_package_metadata(content: bytes) -> dict[str, Any]:
    metadata = {
        "comment_count": 0,
        "revision_count": 0,
        "hidden_text_count": 0,
        "field_count": 0,
    }
    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            names = set(archive.namelist())
            if "word/comments.xml" in names:
                comments_xml = archive.read("word/comments.xml").decode("utf-8", errors="ignore")
                metadata["comment_count"] = comments_xml.count("<w:comment ")
            document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore") if "word/document.xml" in names else ""
            metadata["revision_count"] = len(re.findall(r"<w:(?:ins|del|moveFrom|moveTo)\b", document_xml))
            metadata["hidden_text_count"] = document_xml.count("<w:vanish")
            metadata["field_count"] = document_xml.count("<w:fldChar") + document_xml.count("<w:instrText")
    except Exception:
        return metadata
    return metadata


def _parse_docx_xml_bundle(content: bytes) -> dict[str, Any]:
    try:
        from lxml import etree
    except Exception as exc:
        return {"error": str(exc), "paragraphs": [], "meta": {}}

    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            document_xml = archive.read("word/document.xml")
            root = etree.fromstring(document_xml)
            paragraphs: list[dict[str, Any]] = []
            section_breaks: list[int] = [0]
            for p_idx, p in enumerate(root.xpath(".//w:body/w:p", namespaces=ns)):
                runs = []
                for r_idx, run in enumerate(p.xpath("./w:r", namespaces=ns)):
                    size_val = _first_or_none(run.xpath("./w:rPr/w:sz/@w:val", namespaces=ns))
                    runs.append({
                        "index": r_idx,
                        "text": "".join(str(node) for node in run.xpath(".//w:t/text()", namespaces=ns)),
                        "font_ascii": _first_or_none(run.xpath("./w:rPr/w:rFonts/@w:ascii", namespaces=ns)),
                        "font_east_asia": _first_or_none(run.xpath("./w:rPr/w:rFonts/@w:eastAsia", namespaces=ns)),
                        "font_size_half_points": int(size_val) if str(size_val).isdigit() else None,
                        "bold": bool(run.xpath("./w:rPr/w:b", namespaces=ns)),
                        "italic": bool(run.xpath("./w:rPr/w:i", namespaces=ns)),
                        "superscript": bool(run.xpath("./w:rPr/w:vertAlign[@w:val='superscript']", namespaces=ns)),
                        "subscript": bool(run.xpath("./w:rPr/w:vertAlign[@w:val='subscript']", namespaces=ns)),
                    })
                paragraphs.append({
                    "xml_path": str(PurePosixPath("word/document.xml") / f"p[{p_idx}]"),
                    "style_id": _first_or_none(p.xpath("./w:pPr/w:pStyle/@w:val", namespaces=ns)) or "",
                    "text_runs": runs,
                    "numbering": {
                        "num_id": _first_or_none(p.xpath("./w:pPr/w:numPr/w:numId/@w:val", namespaces=ns)),
                        "ilvl": _first_or_none(p.xpath("./w:pPr/w:numPr/w:ilvl/@w:val", namespaces=ns)),
                    },
                    "page_break_before": bool(p.xpath("./w:pPr/w:pageBreakBefore", namespaces=ns)),
                    "has_page_break": bool(p.xpath(".//w:br[@w:type='page']", namespaces=ns)),
                    "footnote_refs": [str(item) for item in p.xpath(".//w:footnoteReference/@w:id", namespaces=ns)],
                    "endnote_refs": [str(item) for item in p.xpath(".//w:endnoteReference/@w:id", namespaces=ns)],
                    "comment_refs": [str(item) for item in p.xpath(".//w:commentReference/@w:id", namespaces=ns)],
                    "field_codes": [str(item).strip() for item in p.xpath(".//w:instrText/text()", namespaces=ns) if str(item).strip()],
                })
                if p.xpath("./w:pPr/w:sectPr", namespaces=ns):
                    next_index = p_idx + 1
                    if next_index not in section_breaks:
                        section_breaks.append(next_index)
            meta = {
                "paragraph_count": len(paragraphs),
                "has_comments_xml": _zip_exists(archive, "word/comments.xml"),
                "has_footnotes_xml": _zip_exists(archive, "word/footnotes.xml"),
                "has_endnotes_xml": _zip_exists(archive, "word/endnotes.xml"),
                "has_numbering_xml": _zip_exists(archive, "word/numbering.xml"),
            }
            return {"paragraphs": paragraphs, "section_breaks": section_breaks, "meta": meta}
    except Exception as exc:
        return {"error": str(exc), "paragraphs": [], "meta": {}}


def _zip_exists(archive: zipfile.ZipFile, name: str) -> bool:
    try:
        archive.getinfo(name)
        return True
    except KeyError:
        return False


def _first_or_none(values: list[Any]) -> Any:
    return values[0] if values else None


def _extract_citations(text: str) -> list[dict[str, Any]]:
    citations = []
    for m in re.finditer(r"\[(\d+(?:[,,\-–]\d+)*)\]", text):
        raw = m.group(0)
        inner = m.group(1)
        numbers: list[int] = []
        for part in re.split(r"[,,，]", inner):
            part = part.strip()
            if "-" in part or "–" in part:
                try:
                    a_str, b_str = re.split(r"[-–]", part)
                    numbers.extend(range(int(a_str), int(b_str) + 1))
                except ValueError:
                    pass
            else:
                try:
                    numbers.append(int(part))
                except ValueError:
                    pass
        citations.append({"raw": raw, "numbers": numbers, "offset": m.start()})
    return citations


def _heading_level(style_name: str) -> int:
    match = re.search(r"heading\s*(\d+)", style_name, re.I)
    if match:
        return int(match.group(1))
    if "标题" in style_name:
        zh_match = re.search(r"标题\s*(\d+)", style_name)
        if zh_match:
            return int(zh_match.group(1))
        return 1
    return 0


def _to_float(value) -> float | None:
    if value is None:
        return None
    try:
        numeric = float(value)
    except Exception:
        return None
    if abs(numeric) > 1000:
        return round(numeric / 12700, 2)
    return numeric


def _to_pt(value) -> float | None:
    if value is None:
        return None
    try:
        return round(float(value.pt), 2)
    except Exception:
        return None


def _to_cm(value) -> float | None:
    if value is None:
        return None
    try:
        return round(float(value.cm), 2)
    except Exception:
        return None
