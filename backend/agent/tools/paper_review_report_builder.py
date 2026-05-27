"""Build paper review reports in Markdown, DOCX, and PDF formats."""

from __future__ import annotations

from typing import Any


def build_markdown_report(
    *,
    review_output: dict[str, Any],
    evidence_pack: dict[str, Any],
) -> str:
    markdown = str(review_output.get("markdown_report") or "").strip()
    if markdown:
        return markdown

    doc = evidence_pack.get("document") or {}
    issues = list(review_output.get("issues") or evidence_pack.get("issues") or [])
    limitations = list(
        review_output.get("limitations") or evidence_pack.get("limitations") or []
    )
    score = evidence_pack.get("score", 0)
    summary = str(review_output.get("summary") or "")

    high = [i for i in issues if i.get("severity") == "high"]
    medium = [i for i in issues if i.get("severity") == "medium"]
    low = [i for i in issues if i.get("severity") == "low"]

    lines = [
        "# 论文查非与格式审阅辅助报告",
        "",
        "## 一、总体结论",
        "",
        f"- 文档类型：{doc.get('document_type', 'unknown')}",
        f"- 文件名：{doc.get('file_name', 'unknown')}",
        f"- 综合评分：{score} / 100",
        f"- 问题数量：{len(issues)} 个（高 {len(high)} / 中 {len(medium)} / 低 {len(low)}）",
    ]
    if summary:
        lines.append(f"- 总结：{summary}")

    if high:
        lines.extend(["", "## 二、高优先级问题", ""])
        for idx, issue in enumerate(high, start=1):
            lines.extend(_format_issue(idx, issue))

    if medium:
        lines.extend(["", "## 三、中优先级问题", ""])
        for idx, issue in enumerate(medium, start=1):
            lines.extend(_format_issue(idx, issue))

    if low:
        lines.extend(["", "## 四、低优先级问题", ""])
        for idx, issue in enumerate(low, start=1):
            lines.extend(_format_issue(idx, issue))

    need_review = [i for i in issues if i.get("need_human_review")]
    if need_review:
        lines.extend(["", "## 五、需人工复核项", ""])
        for idx, issue in enumerate(need_review, start=1):
            lines.append(f"{idx}. **{issue.get('title', '')}** — {issue.get('location', '')}")

    if limitations:
        lines.extend(["", "## 六、局限性说明", ""])
        for limit in limitations:
            lines.append(f"- {limit}")

    lines.extend([
        "",
        "## 七、修改优先级建议",
        "",
        "1. 先补充结构缺失内容（摘要、关键词、参考文献）。",
        "2. 再统一标题与正文格式（字号、行距、页边距）。",
        "3. 最后处理标点、空格、错别字等细节。",
        "",
        "---",
        "",
        "*本报告为AI辅助生成，仅供修改参考，不代替正式检测或导师审核。*",
    ])

    return "\n".join(lines)


def _format_issue(idx: int, issue: dict[str, Any]) -> list[str]:
    title = issue.get("title") or issue.get("code") or f"问题 {idx}"
    category = issue.get("category", "")
    location = ""
    loc = issue.get("location")
    if isinstance(loc, dict):
        location = str(loc.get("section") or loc.get("paragraph_index") or "")
    evidence = str(issue.get("evidence") or "")
    impact = str(issue.get("impact") or "")
    suggestion = str(issue.get("suggestion") or "")

    lines = [
        f"### {idx}. {title}",
        "",
    ]
    if category:
        lines.append(f"- 类型：{category}")
    if location:
        lines.append(f"- 位置：{location}")
    if evidence:
        lines.append(f"- 证据：{evidence}")
    if impact:
        lines.append(f"- 影响：{impact}")
    if suggestion:
        lines.append(f"- 修改建议：{suggestion}")
    lines.append("")
    return lines


def build_docx_report(markdown_report: str, meta: dict[str, Any] | None = None) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm
    except Exception as exc:
        raise ValueError(f"docx dependency unavailable: {exc}") from exc

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)

    title = doc.add_heading("论文查非与格式审阅辅助报告", level=0)
    title.alignment = 1  # center

    for line in markdown_report.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. "):
            doc.add_paragraph(stripped, style="List Number")
        elif stripped == "---":
            doc.add_paragraph("_" * 40)
        else:
            doc.add_paragraph(stripped)

    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_pdf_report(markdown_report: str, meta: dict[str, Any] | None = None) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except Exception as exc:
        raise ValueError(f"reportlab dependency unavailable: {exc}") from exc

    from io import BytesIO

    buffer = BytesIO()
    doc_template = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    story = []

    for line in markdown_report.splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
            continue
        if stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], styles["Title"]))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], styles["Heading2"]))
        elif stripped.startswith("### "):
            story.append(Paragraph(stripped[4:], styles["Heading3"]))
        elif stripped == "---":
            story.append(Spacer(1, 12))
        else:
            story.append(Paragraph(stripped, styles["Normal"]))

    doc_template.build(story)
    return buffer.getvalue()
