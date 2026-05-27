from __future__ import annotations

from io import BytesIO

from docx import Document

from agent.tools.paper_template_evidence import load_writing_guide_evidence


def _guide_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("写作指南", level=1)
    doc.add_paragraph("摘要应说明研究目的、方法、结果和结论。")
    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


def test_load_writing_guide_evidence_reads_template_from_object_storage(monkeypatch):
    class FakeStorage:
        def get_bytes(self, *, bucket, object_key):
            assert bucket == "paper-templates"
            assert object_key == "cqupt/graduate-thesis/2022/writing-guide.docx"
            return (
                _guide_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    monkeypatch.setattr(
        "app.services.object_storage.factory.build_object_storage",
        lambda: FakeStorage(),
    )

    evidence = load_writing_guide_evidence("cqupt_graduate_thesis_2022")

    assert evidence is not None
    assert evidence["template_id"] == "cqupt_graduate_thesis_2022"
    assert evidence["role"] == "writing_guide"
    assert evidence["bucket"] == "paper-templates"
    assert evidence["object_key"] == "cqupt/graduate-thesis/2022/writing-guide.docx"
    assert "摘要应说明研究目的" in evidence["text"]
    assert evidence["document_type"] == "docx"
