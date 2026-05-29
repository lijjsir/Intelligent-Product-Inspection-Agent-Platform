from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from agent.contracts.quality_contracts import NormalizedAttachment, NormalizedRequest
from agent.router.contracts import AgentPlanStep
from agent.router.manager_loop import ManagerLoop
from agent.router.manager_policy import ManagerPolicy
from agent.router.manager_state import ManagerState
from agent.router.executors.file_executor import FileExecutor


def _docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("论文标题", level=1)
    doc.add_paragraph("这是  一个待检查段落。")
    doc.add_paragraph("参考文献")
    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _ready_runtime_status() -> dict:
    return {
        "ok": True,
        "status": "healthy",
        "engines_used": ["rule", "pycorrector", "macro_correct", "languagetool", "vale"],
        "engine_status": [
            {"name": "pycorrector", "ok": True, "detail": "installed"},
            {"name": "macro_correct", "ok": True, "detail": "installed"},
            {"name": "languagetool", "ok": True, "detail": "http://localhost:8010"},
            {"name": "vale", "ok": True, "detail": "vale 3"},
        ],
    }


@pytest.mark.asyncio
async def test_manager_policy_routes_paper_queries_to_paper_format_check():
    policy = ManagerPolicy()
    request = NormalizedRequest(
        request_id="req-1",
        workflow_run_id="wf-1",
        org_id="org-1",
        user_id="user-1",
        session_id="session-1",
        query="请帮我检查这篇论文的格式和错别字",
        attachments=[NormalizedAttachment(name="paper.docx", kind="file")],
        ext={"surface": "chat"},
    )

    state = policy.initialize_state(request)
    understanding = await policy.understand(state)

    assert understanding.intent == "paper_format_check"
    assert understanding.needs == ["file.paper_format_check", "chat.response.compose"]


@pytest.mark.asyncio
async def test_manager_loop_returns_paper_format_report(monkeypatch):
    stored_objects: dict[tuple[str, str], tuple[bytes, str | None]] = {}
    monkeypatch.setattr("app.services.paper_review_runtime_service.PaperReviewRuntimeService.diagnose_sync", _ready_runtime_status)
    monkeypatch.setattr("agent.tools.paper_format_checker._check_pycorrector", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_macro_correct", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_language_tool", lambda parsed, file_name: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_vale", lambda parsed: [])
    monkeypatch.setattr(
        "app.services.object_storage.resolver.read_attachment_bytes",
        lambda attachment: (_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    )

    output = await ManagerLoop().run(
        NormalizedRequest(
            request_id="req-2",
            workflow_run_id="wf-2",
            org_id="org-1",
            user_id="user-1",
            session_id="session-1",
            query="帮我做论文查非",
            attachments=[
                NormalizedAttachment(
                    name="paper.docx",
                    kind="file",
                    bucket="test",
                    object_key="paper.docx",
                )
            ],
            ext={
                "surface": "chat",
                "template_id": "cqupt_graduate_thesis_2022",
            },
        )
    )

    artifacts = output.agent_output["artifacts"]
    paper_reports = [item for item in artifacts if item["type"] == "paper_format_report"]

    assert output.route_decision.sub_route == "paper_format_check"
    assert output.agent_output["message_type"] == "file_answer"
    assert "已整理出以下修改意见" in output.agent_output["answer"]
    assert paper_reports
    assert paper_reports[0]["content"]["issues"]
    assert paper_reports[0]["content"]["chat_advice"].startswith("已整理出以下修改意见")
    assert len(paper_reports[0]["content"]["issues"]) == paper_reports[0]["content"]["issue_count"]
    assert paper_reports[0]["content"]["report_files"] == []
    assert "issues" in paper_reports[0]["content"]


@pytest.mark.asyncio
async def test_manager_loop_preserves_paper_report_when_composed_response_is_missing(monkeypatch):
    monkeypatch.setattr("app.services.paper_review_runtime_service.PaperReviewRuntimeService.diagnose_sync", _ready_runtime_status)
    monkeypatch.setattr("agent.tools.paper_format_checker._check_pycorrector", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_macro_correct", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_language_tool", lambda parsed, file_name: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_vale", lambda parsed: [])
    monkeypatch.setattr(
        "app.services.object_storage.resolver.read_attachment_bytes",
        lambda attachment: (_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    )

    original_find_composed = ManagerLoop._find_composed

    def fake_find_composed(state):
        return None

    monkeypatch.setattr(ManagerLoop, "_find_composed", staticmethod(fake_find_composed))
    try:
        output = await ManagerLoop().run(
            NormalizedRequest(
                request_id="req-2b",
                workflow_run_id="wf-2b",
                org_id="org-1",
                user_id="user-1",
                session_id="session-1",
                query="帮我做论文查非",
                attachments=[
                    NormalizedAttachment(
                        name="paper.docx",
                        kind="file",
                        bucket="test",
                        object_key="paper.docx",
                    )
                ],
                ext={
                    "surface": "chat",
                    "template_id": "cqupt_graduate_thesis_2022",
                },
            )
        )
    finally:
        monkeypatch.setattr(ManagerLoop, "_find_composed", staticmethod(original_find_composed))

    assert output.agent_output["message_type"] == "file_answer"
    assert output.agent_output["paper_format_report"]["issues"]
    assert output.agent_output["answer"].startswith("已整理出以下修改意见")


@pytest.mark.asyncio
async def test_file_executor_defaults_paper_check_to_cqupt_template(monkeypatch):
    monkeypatch.setattr("app.services.paper_review_runtime_service.PaperReviewRuntimeService.diagnose_sync", _ready_runtime_status)
    monkeypatch.setattr("agent.tools.paper_format_checker._check_pycorrector", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_macro_correct", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_language_tool", lambda parsed, file_name: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_vale", lambda parsed: [])
    monkeypatch.setattr(
        "app.services.object_storage.resolver.read_attachment_bytes",
        lambda attachment: (_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    )

    step = AgentPlanStep(
        step_id="paper-1",
        capability_key="file.paper_format_check",
        agent="chat",
        operation="paper_format_check",
        mode="report",
        input={"attachments": []},
    )
    state = ManagerState(
        request_id="req-default-template",
        workflow_run_id="wf-default-template",
        org_id="org-1",
        user_id="user-1",
        session_id="session-1",
        original_query="paper format check",
        attachments=[
            {
                "name": "paper.docx",
                "kind": "file",
                "bucket": "test",
                "object_key": "paper.docx",
            }
        ],
    )
    request = NormalizedRequest(
        request_id="req-default-template",
        workflow_run_id="wf-default-template",
        org_id="org-1",
        user_id="user-1",
        session_id="session-1",
        query="paper format check",
        attachments=[NormalizedAttachment(name="paper.docx", kind="file")],
        ext={"surface": "chat"},
    )

    _obs, artifacts = await FileExecutor().execute(step, state, request)

    content = artifacts[0].content
    assert content["template_id"] == "cqupt_graduate_thesis_2022"
    assert "issues" in content


@pytest.mark.asyncio
async def test_manager_loop_first_stage_still_returns_file_answer_when_ai_review_would_fail(monkeypatch):
    monkeypatch.setattr("app.services.paper_review_runtime_service.PaperReviewRuntimeService.diagnose_sync", _ready_runtime_status)
    monkeypatch.setattr("agent.tools.paper_format_checker._check_pycorrector", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_macro_correct", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_language_tool", lambda parsed, file_name: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_vale", lambda parsed: [])
    monkeypatch.setattr(
        "app.services.object_storage.resolver.read_attachment_bytes",
        lambda attachment: (_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    )

    async def fake_load_writing_guide_evidence(*args, **kwargs):
        return {"template_id": "cqupt_graduate_thesis_2022", "source": "none", "clauses": []}

    async def fake_generate_ai_review_output(*args, **kwargs):
        from agent.tools.paper_review_ai import PaperReviewModelError

        raise PaperReviewModelError("Ai-Review 模型调用失败：timeout")

    monkeypatch.setattr(
        "agent.tools.paper_template_evidence.load_writing_guide_evidence",
        fake_load_writing_guide_evidence,
    )
    monkeypatch.setattr(
        "agent.tools.paper_review_ai.generate_ai_review_output",
        fake_generate_ai_review_output,
    )

    output = await ManagerLoop().run(
        NormalizedRequest(
            request_id="req-ai-fail",
            workflow_run_id="wf-ai-fail",
            org_id="org-1",
            user_id="user-1",
            session_id="session-1",
            query="帮我做论文查非",
            attachments=[
                NormalizedAttachment(
                    name="paper.docx",
                    kind="file",
                    bucket="test",
                    object_key="paper.docx",
                )
            ],
            ext={"surface": "chat", "template_id": "cqupt_graduate_thesis_2022"},
        ),
        db_session=object(),
    )

    assert output.status == "completed"
    assert output.agent_output["message_type"] == "file_answer"
    assert output.agent_output["answer"]
    assert output.agent_output["paper_format_report"]["report_files"] == []


@pytest.mark.asyncio
async def test_save_report_files_fallback_uses_rule_issues_as_authoritative(monkeypatch):
    stored_objects: dict[tuple[str, str], tuple[bytes, str | None]] = {}

    class FakeStorage:
        def put_bytes(self, *, bucket, object_key, data, content_type=None):
            stored_objects[(bucket, object_key)] = (data, content_type)
            return {
                "bucket": bucket,
                "object_key": object_key,
                "url": f"/api/v1/chat/files/{bucket}/{object_key}",
                "content_type": content_type,
                "size_bytes": len(data),
            }

    monkeypatch.setattr(
        "app.services.object_storage.factory.build_object_storage",
        lambda: FakeStorage(),
    )

    state = ManagerState(
        request_id="req-3",
        workflow_run_id="wf-3",
        org_id="org-1",
        user_id="user-1",
        session_id="session-1",
        assistant_message_id="assistant-1",
        original_query="查非",
    )
    request = NormalizedRequest(
        request_id="req-3",
        workflow_run_id="wf-3",
        org_id="org-1",
        user_id="user-1",
        session_id="session-1",
        query="查非",
        attachments=[],
        ext={},
    )
    merged = {
        "score": 76,
        "summary": "规则检查发现问题。",
        "issues": [
            {
                "code": "structure.abstract_missing",
                "title": "规则缺少摘要",
                "severity": "high",
                "category": "structure",
                "evidence": "未找到摘要",
                "location": {"section": "abstract"},
                "suggestion": "补充摘要",
            }
        ],
        "limitations": [],
        "review_evidence_pack": {
            "score": 76,
            "issues": [],
            "limitations": [],
            "document": {"file_name": "paper.docx", "document_type": "docx"},
        },
        "ai_review_output": {
            "answer": "模型回答",
            "summary": "模型总结",
            "markdown_report": "",
            "issues": [
                {
                    "title": "模型不应覆盖规则问题",
                    "severity": "low",
                    "category": "text",
                    "location": "model",
                    "evidence": "model-only issue",
                    "suggestion": "不应进入最终问题来源",
                }
            ],
            "model_used": True,
        },
    }

    report_files = await FileExecutor._save_report_files(
        merged=merged,
        state=state,
        request=request,
    )

    markdown_file = next(item for item in report_files if item["format"] == "md")
    assert markdown_file["url"].startswith("/api/v1/chat/files/chat-reports/")
    assert any(item["format"] == "docx" for item in report_files)
    markdown_bytes = stored_objects[(markdown_file["bucket"], markdown_file["object_key"])][0]
    markdown = markdown_bytes.decode("utf-8")
    assert "规则缺少摘要" in markdown
    assert "模型不应覆盖规则问题" not in markdown
    assert merged["ai_review_output"]["issues"][0]["code"] == "structure.abstract_missing"


def test_report_file_url_uses_backend_download_route_for_storage_objects():
    url = FileExecutor._report_file_url(
        {"url": "http://127.0.0.1:9000/chat-reports/org/user/report.md"},
        "org/user/report.md",
    )

    assert url == "/api/v1/chat/files/chat-reports/org/user/report.md"
