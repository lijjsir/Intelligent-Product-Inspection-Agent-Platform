from __future__ import annotations

import sys
import types
from io import BytesIO

from docx import Document

from agent.tools.file_parsers import parse_docx_bytes, parse_pdf_bytes, parse_tex_bytes
from agent.tools.paper_format_checker import _check_pycorrector, check_paper_format
from agent.tools.paper_format_templates import get_paper_template


def _ready_runtime_status() -> dict:
    return {
        "ok": True,
        "status": "healthy",
        "engines_used": ["rule", "pycorrector", "macro_correct", "languagetool", "vale"],
        "engine_status": [
            {"name": "docx", "ok": True, "detail": "installed"},
            {"name": "lxml", "ok": True, "detail": "installed"},
            {"name": "pycorrector", "ok": True, "detail": "installed"},
            {"name": "macro_correct", "ok": True, "detail": "installed"},
            {"name": "languagetool", "ok": True, "detail": "http://localhost:8010"},
            {"name": "vale", "ok": True, "detail": "vale 3"},
        ],
    }


def _build_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("论文标题。", level=1)
    doc.add_paragraph("这是  一个包含连续空格的正文段落，mix punctuation, 需要检查。")
    doc.add_paragraph("关键词：查非；格式检查")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_parse_docx_bytes_extracts_structure_and_layout():
    parsed = parse_docx_bytes(_build_docx_bytes())

    assert parsed["kind"] == "docx"
    assert parsed["paragraphs"]
    assert parsed["headings"][0]["text"].startswith("论文标题")
    assert "page_layout" in parsed
    assert parsed["section_count"] >= 1


def test_parse_docx_enhanced_extracts_template_check_metadata():
    from agent.tools.paper_docx_parser import parse_docx_enhanced

    parsed = parse_docx_enhanced(_build_docx_bytes())

    assert "orientation" in parsed["page_layout"]
    assert "gutter_cm" in parsed["page_layout"]
    assert "toc_entries" in parsed
    assert "table_titles" in parsed
    assert "formula_numbers" in parsed
    assert "word_metadata" in parsed
    assert "ooxml" in parsed
    assert "style_id" in parsed["paragraphs"][0]
    assert "xml_path" in parsed["paragraphs"][0]
    assert "text_runs" in parsed["paragraphs"][0]
    assert "word_section_index" in parsed["paragraphs"][0]


def test_parse_tex_bytes_extracts_commands_and_sections():
    content = r"""
\documentclass{article}
\title{A Demo Paper}
\author{Tester}
\begin{document}
\maketitle
\section{引言}
这是正文。
\begin{figure}
\end{figure}
\end{document}
""".encode("utf-8")
    parsed = parse_tex_bytes(content)

    assert parsed["kind"] == "tex"
    assert parsed["commands"]["title"] == "A Demo Paper"
    assert parsed["sections"][0]["title"] == "引言"
    assert parsed["figure_count"] == 1


def test_parse_pdf_bytes_extracts_page_evidence(monkeypatch):
    class FakePage:
        def __init__(self, text: str):
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class FakePdfReader:
        def __init__(self, stream):
            self.pages = [
                FakePage("1 Introduction\nThis is page one."),
                FakePage("References\n[1] Example source."),
            ]

    monkeypatch.setitem(sys.modules, "pypdf", types.SimpleNamespace(PdfReader=FakePdfReader))

    parsed = parse_pdf_bytes(b"%PDF-1.7")

    assert parsed["kind"] == "pdf"
    assert parsed["page_count"] == 2
    assert parsed["pages"][0]["page_no"] == 1
    assert parsed["pages"][0]["blocks"][0]["type"] == "text"
    assert parsed["headings"][0]["text"] == "1 Introduction"
    assert parsed["references"] == ["[1] Example source."]
    assert parsed["layout"]["analysis"] == "text_only"


def test_check_paper_format_reports_missing_sections_for_docx(monkeypatch):
    monkeypatch.setattr("app.services.paper_review_runtime_service.PaperReviewRuntimeService.diagnose_sync", _ready_runtime_status)
    monkeypatch.setattr("agent.tools.paper_format_checker._check_pycorrector", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_macro_correct", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_language_tool", lambda parsed, file_name: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_vale", lambda parsed: [])
    report = check_paper_format(
        parsed=parse_docx_bytes(_build_docx_bytes()),
        file_name="paper.docx",
        query="帮我查非",
        template_id=None,
    )

    issue_codes = {item["code"] for item in report["issues"]}
    assert report["document_type"] == "docx"
    assert report["template_id"] == "generic_cn_thesis"
    assert "structure.abstract_missing" in issue_codes
    assert "text.multiple_spaces" in issue_codes
    assert report["limitations"]
    mixed_issue = next(item for item in report["issues"] if item["code"] == "text.mixed_punctuation")
    assert mixed_issue["evidence"] != "文本中同时出现中文和英文标点组合"
    assert "mix punctuation," in mixed_issue["evidence"]
    assert mixed_issue["location"]["display_text"]
    heading_issue = next(item for item in report["issues"] if item["code"] == "text.heading_trailing_punct")
    assert heading_issue["evidence"] == "论文标题。"
    assert heading_issue["location"]["display_text"]


def test_check_paper_format_reports_fullwidth_ascii_with_exact_location(monkeypatch):
    monkeypatch.setattr("app.services.paper_review_runtime_service.PaperReviewRuntimeService.diagnose_sync", _ready_runtime_status)
    monkeypatch.setattr("agent.tools.paper_format_checker._check_pycorrector", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_macro_correct", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_language_tool", lambda parsed, file_name: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_vale", lambda parsed: [])
    parsed = {
        "kind": "docx",
        "text": "摘要\n这是ＡBC格式错误。\n关键词：查非",
        "headings": [
            {"text": "摘要", "level": 1, "paragraph_index": 0, "section_index": 1},
        ],
        "paragraphs": [
            {"index": 0, "text": "摘要", "heading_level": 1, "section_title": "摘要", "section_index": 1, "paragraph_no": 0},
            {"index": 1, "text": "这是ＡBC格式错误。", "heading_level": 0, "section_title": "摘要", "section_index": 1, "paragraph_no": 1},
            {"index": 2, "text": "关键词：查非", "heading_level": 0, "section_title": "摘要", "section_index": 1, "paragraph_no": 2},
        ],
        "page_layout": {},
    }

    report = check_paper_format(
        parsed=parsed,
        file_name="paper.docx",
        query="帮我查非",
        template_id=None,
    )

    fullwidth_issue = next(item for item in report["issues"] if item["code"] == "text.fullwidth_ascii")
    assert "这是ＡBC格式错误。" == fullwidth_issue["evidence"]
    assert fullwidth_issue["location"]["display_text"] == "第1节《摘要》下第1段"


def test_check_paper_format_reports_tex_limitations():
    parsed = parse_tex_bytes(
        r"""
\documentclass{article}
\title{A Demo Paper}
\author{Tester}
\begin{document}
\maketitle
\section{引言}
\end{document}
""".encode("utf-8")
    )
    report = check_paper_format(
        parsed=parsed,
        file_name="paper.tex",
        query="检查论文格式",
        template_id="generic_cn_thesis",
    )

    issue_codes = {item["code"] for item in report["issues"]}
    assert report["document_type"] == "tex"
    assert "tex.abstract_missing" in issue_codes
    assert any("LaTeX" in item for item in report["limitations"])


def test_check_macro_correct_uses_batch_detectors(monkeypatch):
    from agent.tools.paper_format_checker import _check_macro_correct

    monkeypatch.setattr("app.core.config.settings.paper_check_macro_correct_enabled", True)
    monkeypatch.setattr(
        "agent.tools.paper_format_checker.run_macro_correct_token",
        lambda texts: (
            [{"errors": [{"wrong": "錯", "right": "错", "begin": 2, "end": 3, "rule_id": "token"}]}],
            "macro_correct.token",
        ),
    )
    monkeypatch.setattr(
        "agent.tools.paper_format_checker.run_macro_correct_punct",
        lambda texts: (
            [{"errors": [{"wrong": ",", "right": "，", "begin": 5, "end": 6, "rule_id": "punct"}]}],
            "macro_correct.punct",
        ),
    )

    parsed = {
        "paragraphs": [
            {
                "index": 0,
                "text": "这是錯字,示例。",
                "section_title": "摘要",
                "section_index": 1,
                "paragraph_no": 1,
            }
        ]
    }

    issues = _check_macro_correct(parsed)

    assert len(issues) == 2
    assert {item.engine_rule_id for item in issues} == {"macro_correct.token", "macro_correct.punct"}
    assert any(item.actual == {"wrong": "錯", "right": "错"} for item in issues)
    assert any(item.actual == {"wrong": ",", "right": "，"} for item in issues)


def test_check_paper_format_returns_runtime_not_ready_for_docx(monkeypatch):
    monkeypatch.setattr(
        "app.services.paper_review_runtime_service.PaperReviewRuntimeService.diagnose_sync",
        lambda: {
            "ok": False,
            "status": "unhealthy",
            "engines_used": ["rule", "pycorrector", "macro_correct", "languagetool", "vale"],
            "engine_status": [
                {"name": "pycorrector", "ok": False, "detail": "missing"},
            ],
        },
    )

    report = check_paper_format(
        parsed=parse_docx_bytes(_build_docx_bytes()),
        file_name="paper.docx",
        query="帮我查非",
        template_id=None,
    )

    assert report["runtime_ready"] is False
    assert report["issues"] == []
    assert "论文检测环境未就绪" in report["summary"]
    assert report["engine_status"][0]["ok"] is False


def test_check_pycorrector_supports_current_class_based_api(monkeypatch):
    class FakeCorrector:
        def correct(self, text: str):
            assert text == "这是測试段落。"
            return {
                "source": text,
                "target": "这是测试段落。",
                "errors": [("測", "测", 2)],
            }

    monkeypatch.setattr(
        "agent.tools.paper_review_pycorrector._build_corrector",
        lambda: (FakeCorrector().correct, "pycorrector.Corrector.correct"),
    )

    issues = _check_pycorrector(
        {
            "paragraphs": [
                {
                    "index": 0,
                    "paragraph_no": 1,
                    "section_title": "引言",
                    "section_index": 1,
                    "heading_level": 0,
                    "text": "这是測试段落。",
                }
            ]
        }
    )

    assert len(issues) == 1
    assert issues[0].engine == "pycorrector"
    assert issues[0].engine_rule_id == "pycorrector.Corrector.correct"


def test_check_pycorrector_chunks_document_and_maps_back_to_paragraph(monkeypatch):
    monkeypatch.setattr("agent.tools.paper_format_checker.settings.paper_check_pycorrector_chunk_chars", 20)
    monkeypatch.setattr("agent.tools.paper_format_checker.settings.paper_check_pycorrector_timeout_sec", 5)

    def fake_run_pycorrector(text: str):
        if "第二段包含測试问题。" in text:
            begin = text.index("測")
            return ({"errors": [("測", "测", begin)]}, "pycorrector.Corrector.correct")
        return ({"errors": []}, "pycorrector.Corrector.correct")

    monkeypatch.setattr("agent.tools.paper_format_checker.run_pycorrector", fake_run_pycorrector)

    issues = _check_pycorrector(
        {
            "paragraphs": [
                {
                    "index": 0,
                    "paragraph_no": 1,
                    "section_title": "引言",
                    "section_index": 1,
                    "heading_level": 0,
                    "text": "第一段内容较短。",
                },
                {
                    "index": 1,
                    "paragraph_no": 2,
                    "section_title": "引言",
                    "section_index": 1,
                    "heading_level": 0,
                    "text": "第二段包含測试问题。",
                },
            ]
        }
    )

    assert len(issues) == 1
    assert issues[0].location["display_text"] == "第1节《引言》下第2段"
    assert "第二段包含測试问题。" == issues[0].evidence
    assert issues[0].actual == {"wrong": "測", "right": "测"}
    assert "測试" in issues[0].evidence


def test_cqupt_template_is_registered_with_storage_objects():
    template = get_paper_template("cqupt_graduate_thesis_2022")

    assert template["template_id"] == "cqupt_graduate_thesis_2022"
    assert template["name"] == "重庆邮电大学研究生学位论文模板（2022版）"
    assert template["storage"]["bucket"] == "paper-templates"
    object_keys = {item["object_key"] for item in template["storage"]["files"]}
    assert "cqupt/graduate-thesis/2022/word-commented-template.docx" in object_keys
    assert "cqupt/graduate-thesis/2022/writing-guide.docx" in object_keys


def test_check_paper_format_reports_cqupt_template_differences(monkeypatch):
    monkeypatch.setattr("app.services.paper_review_runtime_service.PaperReviewRuntimeService.diagnose_sync", _ready_runtime_status)
    monkeypatch.setattr("agent.tools.paper_format_checker._check_pycorrector", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_macro_correct", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_language_tool", lambda parsed, file_name: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_vale", lambda parsed: [])
    parsed = {
        "kind": "docx",
        "text": "摘要\n关键词：查非\n目录\n引言\n正文段落\n参考文献",
        "headings": [
            {"text": "引言", "level": 1, "paragraph_index": 3, "section_index": 1},
            {"text": "参考文献", "level": 1, "paragraph_index": 5, "section_index": 2},
        ],
        "paragraphs": [
            {
                "index": 0,
                "text": "摘要",
                "heading_level": 1,
                "font_name": "宋体",
                "font_size_pt": 12,
                "line_spacing": 1.5,
                "first_line_indent_pt": 0,
                "section_title": "摘要",
                "section_index": 0,
                "paragraph_no": 0,
            },
            {
                "index": 1,
                "text": "关键词：查非",
                "heading_level": 0,
                "font_name": "宋体",
                "font_size_pt": 12,
                "line_spacing": 1.5,
                "first_line_indent_pt": 0,
                "section_title": "摘要",
                "section_index": 0,
                "paragraph_no": 1,
            },
            {
                "index": 2,
                "text": "目录",
                "heading_level": 1,
                "font_name": "宋体",
                "font_size_pt": 12,
                "line_spacing": 1.5,
                "first_line_indent_pt": 0,
                "section_title": "目录",
                "section_index": 0,
                "paragraph_no": 0,
            },
            {
                "index": 3,
                "text": "引言",
                "heading_level": 1,
                "font_name": "黑体",
                "font_size_pt": 12,
                "line_spacing": 1.0,
                "first_line_indent_pt": 0,
                "section_title": "引言",
                "section_index": 1,
                "paragraph_no": 0,
            },
            {
                "index": 4,
                "text": "正文段落",
                "heading_level": 0,
                "font_name": "Arial",
                "font_size_pt": 10.5,
                "line_spacing": 1.0,
                "first_line_indent_pt": 0,
                "section_title": "引言",
                "section_index": 1,
                "paragraph_no": 1,
            },
            {
                "index": 5,
                "text": "参考文献",
                "heading_level": 1,
                "font_name": "黑体",
                "font_size_pt": 12,
                "line_spacing": 1.5,
                "first_line_indent_pt": 0,
                "section_title": "参考文献",
                "section_index": 2,
                "paragraph_no": 0,
            }
        ],
        "page_layout": {
            "top_margin_cm": 2.54,
            "bottom_margin_cm": 2.54,
            "left_margin_cm": 3.18,
            "right_margin_cm": 3.18,
        },
    }

    report = check_paper_format(
        parsed=parsed,
        file_name="paper.docx",
        query="按重庆邮电大学模板查非",
        template_id="cqupt_graduate_thesis_2022",
    )

    issue_codes = {item["code"] for item in report["issues"]}
    assert report["template_id"] == "cqupt_graduate_thesis_2022"
    assert "template.required_section_missing" in issue_codes
    missing_titles = {item["title"] for item in report["issues"] if item["code"] == "template.required_section_missing"}
    assert "模板要求章节缺失：正文" not in missing_titles
    assert "template.margin_mismatch" in issue_codes
    assert "template.body_font_mismatch" in issue_codes
    assert any("重庆邮电大学" in item for item in report["limitations"])
    font_issue = next(item for item in report["issues"] if item["code"] == "template.body_font_mismatch")
    assert font_issue["location"]["section_title"] == "引言"
    assert "正文样式汇总" in font_issue["location"]["display_text"]
    assert font_issue["actual"]["checked_count"] >= font_issue["actual"]["mismatch_count"] >= 1
    assert font_issue["actual"]["samples"][0]["display_text"] == "第1节《引言》下第1段"
    assert font_issue["actual"]["samples"][0]["text"] == "本研究围绕智能检测平台的关键问题展开分析。"
    assert "本研究围绕智能检测平台的关键问题展开分析。" in font_issue["evidence"]


def test_check_paper_format_reports_missing_body_when_no_main_text(monkeypatch):
    monkeypatch.setattr("app.services.paper_review_runtime_service.PaperReviewRuntimeService.diagnose_sync", _ready_runtime_status)
    monkeypatch.setattr("agent.tools.paper_format_checker._check_pycorrector", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_macro_correct", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_language_tool", lambda parsed, file_name: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_vale", lambda parsed: [])
    parsed = {
        "kind": "docx",
        "text": "摘要\n关键词：查非\n目录\n参考文献\n致谢",
        "headings": [
            {"text": "摘要", "level": 1, "paragraph_index": 0, "section_index": 0},
            {"text": "目录", "level": 1, "paragraph_index": 2, "section_index": 0},
            {"text": "参考文献", "level": 1, "paragraph_index": 3, "section_index": 1},
        ],
        "paragraphs": [
            {"index": 0, "text": "摘要", "heading_level": 1, "section_title": "摘要", "section_index": 0, "paragraph_no": 0},
            {"index": 1, "text": "关键词：查非", "heading_level": 0, "section_title": "摘要", "section_index": 0, "paragraph_no": 1},
            {"index": 2, "text": "目录", "heading_level": 1, "section_title": "目录", "section_index": 0, "paragraph_no": 0},
            {"index": 3, "text": "参考文献", "heading_level": 1, "section_title": "参考文献", "section_index": 1, "paragraph_no": 0},
            {"index": 4, "text": "致谢", "heading_level": 1, "section_title": "致谢", "section_index": 2, "paragraph_no": 0},
        ],
        "page_layout": {},
    }

    report = check_paper_format(
        parsed=parsed,
        file_name="paper.docx",
        query="按重庆邮电大学模板查非",
        template_id="cqupt_graduate_thesis_2022",
    )

    missing_titles = {item["title"] for item in report["issues"] if item["code"] == "template.required_section_missing"}
    assert "模板要求章节缺失：正文" in missing_titles


def test_check_paper_format_covers_checklist_layout_front_matter_and_word_artifacts(monkeypatch):
    monkeypatch.setattr("app.services.paper_review_runtime_service.PaperReviewRuntimeService.diagnose_sync", _ready_runtime_status)
    monkeypatch.setattr("agent.tools.paper_format_checker._check_pycorrector", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_macro_correct", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_language_tool", lambda parsed, file_name: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_vale", lambda parsed: [])
    parsed = {
        "kind": "docx",
        "text": (
            "封面\n摘要\n摘要正文引用了图1和[3]。\n关键词：一个，两个\nABSTRACT\nKeywords: one, two\n"
            "目录\n1 绪论\n1.2 研究背景\n正文见图2-1、表2-1和式（2-1）。\n"
            "图2-2 错位图题\n表2-2 错位表题\nE=mc^2 （2-3）\n"
            "参考文献\n[1] 张三. 论文题名[J]. 期刊, 2020.\n[3] 缺少编号二的文献\n"
        ),
        "headings": [
            {"text": "摘要", "level": 1, "paragraph_index": 1, "section_index": 1},
            {"text": "ABSTRACT", "level": 1, "paragraph_index": 4, "section_index": 2},
            {"text": "目录", "level": 1, "paragraph_index": 6, "section_index": 3},
            {"text": "1 绪论", "level": 1, "paragraph_index": 7, "section_index": 4},
            {"text": "1.2 研究背景", "level": 2, "paragraph_index": 8, "section_index": 4},
            {"text": "参考文献", "level": 1, "paragraph_index": 13, "section_index": 5},
        ],
        "paragraphs": [
            {"index": 0, "text": "封面", "heading_level": 0, "section_title": "", "section_index": 0, "paragraph_no": 1},
            {"index": 1, "text": "摘要", "heading_level": 1, "section_title": "摘要", "section_index": 1, "paragraph_no": 0},
            {"index": 2, "text": "摘要正文引用了图1和[3]。", "heading_level": 0, "section_title": "摘要", "section_index": 1, "paragraph_no": 1, "font_name": "宋体", "font_size_pt": 12, "line_spacing": 1.5, "first_line_indent_pt": 0},
            {"index": 3, "text": "关键词：一个，两个", "heading_level": 0, "section_title": "摘要", "section_index": 1, "paragraph_no": 2, "font_name": "宋体", "font_size_pt": 12, "line_spacing": 1.5},
            {"index": 4, "text": "ABSTRACT", "heading_level": 1, "section_title": "ABSTRACT", "section_index": 2, "paragraph_no": 0},
            {"index": 5, "text": "Keywords: one, two", "heading_level": 0, "section_title": "ABSTRACT", "section_index": 2, "paragraph_no": 1, "font_name": "Times New Roman", "font_size_pt": 12, "line_spacing": 1.5},
            {"index": 6, "text": "目录", "heading_level": 1, "section_title": "目录", "section_index": 3, "paragraph_no": 0},
            {"index": 7, "text": "1 绪论", "heading_level": 1, "section_title": "1 绪论", "section_index": 4, "paragraph_no": 0, "font_name": "黑体", "font_size_pt": 15},
            {"index": 8, "text": "1.2 研究背景", "heading_level": 2, "section_title": "1 绪论", "section_index": 4, "paragraph_no": 0, "font_name": "黑体", "font_size_pt": 14},
            {"index": 9, "text": "正文见图2-1、表2-1和式（2-1）。", "heading_level": 0, "section_title": "1 绪论", "section_index": 4, "paragraph_no": 1, "font_name": "宋体", "font_size_pt": 12, "line_spacing": 1.0, "first_line_indent_pt": 0},
            {"index": 10, "text": "图2-2 错位图题", "heading_level": 0, "section_title": "1 绪论", "section_index": 4, "paragraph_no": 2},
            {"index": 11, "text": "表2-2 错位表题", "heading_level": 0, "section_title": "1 绪论", "section_index": 4, "paragraph_no": 3},
            {"index": 12, "text": "E=mc^2 （2-3）", "heading_level": 0, "section_title": "1 绪论", "section_index": 4, "paragraph_no": 4},
            {"index": 13, "text": "参考文献", "heading_level": 1, "section_title": "参考文献", "section_index": 5, "paragraph_no": 0},
            {"index": 14, "text": "[1] 张三. 论文题名[J]. 期刊, 2020.", "heading_level": 0, "section_title": "参考文献", "section_index": 5, "paragraph_no": 1},
            {"index": 15, "text": "[3] 缺少编号二的文献", "heading_level": 0, "section_title": "参考文献", "section_index": 5, "paragraph_no": 2},
        ],
        "toc_entries": [{"title": "1 绪论", "level": 1, "page": "1"}],
        "figure_titles": ["图2-2 错位图题"],
        "table_titles": ["表2-2 错位表题"],
        "formula_numbers": ["2-3"],
        "references": ["[1] 张三. 论文题名[J]. 期刊, 2020.", "[3] 缺少编号二的文献"],
        "word_metadata": {"comment_count": 1, "revision_count": 2, "hidden_text_count": 1},
        "sections": [
            {
                "page_width_cm": 20.0,
                "page_height_cm": 29.7,
                "orientation": "portrait",
                "top_margin_cm": 2.0,
                "bottom_margin_cm": 3.0,
                "left_margin_cm": 2.0,
                "right_margin_cm": 3.0,
                "gutter_cm": 0.5,
                "header_distance_cm": 1.0,
                "footer_distance_cm": 1.0,
                "header_text": "错误页眉",
                "footer_text": "第 1 页",
            }
        ],
        "page_layout": {
            "page_width_cm": 20.0,
            "page_height_cm": 29.7,
            "orientation": "portrait",
            "top_margin_cm": 2.0,
            "bottom_margin_cm": 3.0,
            "left_margin_cm": 2.0,
            "right_margin_cm": 3.0,
            "gutter_cm": 0.5,
            "header_distance_cm": 1.0,
            "footer_distance_cm": 1.0,
            "header_text": "错误页眉",
            "footer_text": "第 1 页",
        },
    }

    report = check_paper_format(
        parsed=parsed,
        file_name="paper.docx",
        query="按清单和重邮模板查非",
        template_id="cqupt_graduate_thesis_2022",
    )

    issue_codes = {item["code"] for item in report["issues"]}
    assert "template.page_size_mismatch" in issue_codes
    assert "template.gutter_mismatch" in issue_codes
    assert "template.header_footer_mismatch" in issue_codes
    assert "abstract.keyword_count_out_of_range" in issue_codes
    assert "toc.required_entry_missing" in issue_codes
    assert "heading.numbering_discontinuous" in issue_codes
    assert "style.paragraph_indent_missing" in issue_codes
    assert "style.line_spacing_small" in issue_codes
    assert "figure.referenced_caption_missing" in issue_codes
    assert "table.referenced_caption_missing" in issue_codes
    assert "formula.numbering_discontinuous" in issue_codes
    assert "references.numbering_discontinuous" in issue_codes
    assert "references.entry_format_incomplete" in issue_codes
    assert "word.comments_or_revisions_present" in issue_codes


def test_check_paper_format_supports_pdf_text_phase_with_limitations():
    parsed = {
        "kind": "pdf",
        "text": "论文标题\n关键词：查非\n参考文献\n正文中有  连续空格。",
        "page_count": 3,
        "pages": [
            {"page_no": 1, "text": "论文标题"},
            {"page_no": 2, "text": "正文中有  连续空格。"},
        ],
        "headings": [],
        "references": [],
        "layout": {},
    }

    report = check_paper_format(
        parsed=parsed,
        file_name="paper.pdf",
        query="检查论文格式",
        template_id="generic_cn_thesis",
    )

    issue_codes = {item["code"] for item in report["issues"]}
    assert report["document_type"] == "pdf"
    assert "unsupported.document_type" not in issue_codes
    assert "structure.abstract_missing" in issue_codes
    assert "text.multiple_spaces" in issue_codes
    assert any("PDF" in item and "文本抽取" in item for item in report["limitations"])


def test_header_content_check_starts_from_body_section(monkeypatch):
    monkeypatch.setattr("app.services.paper_review_runtime_service.PaperReviewRuntimeService.diagnose_sync", _ready_runtime_status)
    monkeypatch.setattr("agent.tools.paper_format_checker._check_pycorrector", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_macro_correct", lambda parsed: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_language_tool", lambda parsed, file_name: [])
    monkeypatch.setattr("agent.tools.paper_format_checker._check_vale", lambda parsed: [])
    parsed = {
        "kind": "docx",
        "text": "摘要\n关键词：查非\n目录\n1 绪论\n正文段落\n参考文献",
        "headings": [
            {"text": "摘要", "level": 1, "paragraph_index": 0, "section_index": 1},
            {"text": "目录", "level": 1, "paragraph_index": 2, "section_index": 2},
            {"text": "1 绪论", "level": 1, "paragraph_index": 3, "section_index": 3},
            {"text": "参考文献", "level": 1, "paragraph_index": 5, "section_index": 4},
        ],
        "paragraphs": [
            {"index": 0, "text": "摘要", "heading_level": 1, "section_title": "摘要", "section_index": 1, "word_section_index": 0, "paragraph_no": 0},
            {"index": 1, "text": "关键词：查非", "heading_level": 0, "section_title": "摘要", "section_index": 1, "word_section_index": 0, "paragraph_no": 1},
            {"index": 2, "text": "目录", "heading_level": 1, "section_title": "目录", "section_index": 2, "word_section_index": 1, "paragraph_no": 0},
            {"index": 3, "text": "1 绪论", "heading_level": 1, "section_title": "1 绪论", "section_index": 3, "word_section_index": 2, "paragraph_no": 0},
            {"index": 4, "text": "正文段落", "heading_level": 0, "section_title": "1 绪论", "section_index": 3, "word_section_index": 2, "paragraph_no": 1},
            {"index": 5, "text": "参考文献", "heading_level": 1, "section_title": "参考文献", "section_index": 4, "word_section_index": 2, "paragraph_no": 0},
        ],
        "sections": [
            {"header_text": "错误封面页眉", "footer_text": ""},
            {"header_text": "错误前置页眉", "footer_text": ""},
            {"header_text": "重庆邮电大学硕士学位论文", "footer_text": ""},
        ],
        "page_layout": {
            "page_width_cm": 21.0,
            "page_height_cm": 29.7,
            "orientation": "portrait",
            "top_margin_cm": 3.0,
            "bottom_margin_cm": 3.0,
            "left_margin_cm": 3.0,
            "right_margin_cm": 3.0,
            "gutter_cm": 0.0,
            "header_distance_cm": 2.0,
            "footer_distance_cm": 2.0,
        },
    }

    report = check_paper_format(
        parsed=parsed,
        file_name="paper.docx",
        query="按重庆邮电大学模板查非",
        template_id="cqupt_graduate_thesis_2022",
    )

    issue_codes = {item["code"] for item in report["issues"]}
    assert "template.header_content_mismatch" not in issue_codes
